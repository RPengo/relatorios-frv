from pathlib import Path
import json, subprocess, sys
from docx import Document

base=Path('tmp_test'); base.mkdir(exist_ok=True)
docx=base/'in.docx'
d=Document(); d.add_paragraph('RESULTADOS E DISCUSSÃO'); d.add_paragraph('texto antigo'); d.add_paragraph('Tabela 5. Altura de plantas'); d.add_table(rows=1, cols=1); d.add_paragraph('Tabela 6. Produtividade'); d.add_table(rows=1, cols=1); d.add_paragraph('Figura 2. Gráfico'); d.add_paragraph('CONSIDERAÇÕES'); d.add_paragraph('antigo'); d.add_paragraph('REFERÊNCIAS BIBLIOGRÁFICAS'); d.save(docx)
(base/'resultados_discussao.md').write_text('fallback', encoding='utf-8')
(base/'consideracoes.md').write_text('Conforme as condições em que este ensaio foi realizado podemos concluir que:\n\n- a\n- b\n- c\n', encoding='utf-8')
(base/'resultados_discussao_blocos.json').write_text(json.dumps({'blocos_resultados':[{'tipo_anchor':'tabela','numero_anchor':5,'titulo_detectado':'Tabela 5. Altura de plantas','texto':'Texto T5'},{'tipo_anchor':'tabela','numero_anchor':6,'titulo_detectado':'Tabela 6. Produtividade','texto':'Texto T6'}]}, ensure_ascii=False), encoding='utf-8')
subprocess.check_call([sys.executable,'scripts/06_inserir_no_docx.py','--docx-original',str(docx),'--resultados',str(base/'resultados_discussao.md'),'--consideracoes',str(base/'consideracoes.md'),'--blocos-resultados',str(base/'resultados_discussao_blocos.json'),'--saida-dir',str(base)])
out=Document(base/'in_final.docx')
texts=[p.text for p in out.paragraphs if p.text.strip()]
assert 'Tabela 5. Altura de plantas' in texts
assert 'Tabela 6. Produtividade' in texts
assert texts.index('Texto T5') < texts.index('Tabela 5. Altura de plantas')
assert texts.index('Texto T6') < texts.index('Tabela 6. Produtividade')
assert len(out.tables)==2
print('ok')
