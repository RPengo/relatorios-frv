#!/usr/bin/env python3
from __future__ import annotations
import argparse, json, re, sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
ABERTURA_CONSIDERACOES = "Conforme as condições em que este ensaio foi realizado podemos concluir que:"
STYLE_DISC = "FRV_AI_Discussion"
STYLE_CONS = "FRV_AI_Consideracao"

def parse_args():
 p=argparse.ArgumentParser(); p.add_argument('--docx-original',type=Path,required=True); p.add_argument('--resultados',type=Path,default=Path('saida/textos_gerados/resultados_discussao.md')); p.add_argument('--consideracoes',type=Path,default=Path('saida/textos_gerados/consideracoes.md')); p.add_argument('--blocos-resultados',type=Path,default=Path('saida/textos_gerados/resultados_discussao_blocos.json')); p.add_argument('--saida-dir',type=Path,default=Path('saida/relatorios_finalizados')); return p.parse_args()

def norm(t:str)->str: return re.sub(r'\s+',' ',t or '').strip().upper()
def carregar(c:Path)->str: return c.read_text(encoding='utf-8').strip()
def blocos(t:str)->list[str]: return [b.strip() for b in re.split(r'\n\s*\n',t) if b.strip()]

def get_title(doc,titulo):
 a=norm(titulo)
 for p in doc.paragraphs:
  if norm(p.text)==a:return p
 raise RuntimeError(f"Seção '{titulo}' não encontrada")

def ensure_style(doc, name):
 styles=doc.styles
 if name in [s.name for s in styles]: return name
 base=styles['Normal'] if 'Normal' in [s.name for s in styles] else styles[0]
 st=styles.add_style(name,1); st.base_style=base; return name

def formatar(p):
 p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; f=p.paragraph_format; f.line_spacing=1.5; f.first_line_indent=Pt(24)

def insert_before(doc, anchor_p, texto, estilo):
 novo=doc.add_paragraph(texto, style=estilo); formatar(novo)
 parent=anchor_p._element.getparent(); parent.insert(parent.index(anchor_p._element), novo._element)

def remover_estilo_entre(doc, ini, fim, estilos):
 body=doc._body._element; elems=list(body); i1=elems.index(ini._element); i2=elems.index(fim._element)
 for el in elems[i1+1:i2]:
  if isinstance(el, CT_P):
   p_txt=''.join(el.itertext()).strip()
   if re.match(r'^(Tabela|Figura|Quadro)\s+\d+', p_txt, re.IGNORECASE):
    continue
   p_obj=next((p for p in doc.paragraphs if p._element is el), None)
   if p_obj and p_obj.style and p_obj.style.name in estilos: el.getparent().remove(el)

def substituir_resultados(doc, texto_md, json_blocos):
 titulo=get_title(doc,'RESULTADOS E DISCUSSÃO'); cons=get_title(doc,'CONSIDERAÇÕES'); ensure_style(doc, STYLE_DISC)
 remover_estilo_entre(doc, titulo, cons, {STYLE_DISC})
 body=doc._body._element; elems=list(body); i1=elems.index(titulo._element); i2=elems.index(cons._element)
 if json_blocos and json_blocos.exists():
  data=json.loads(json_blocos.read_text(encoding='utf-8')); blocos_res=data.get('blocos_resultados',[])
  for bloco in blocos_res:
   padrao=rf"^{bloco.get('tipo_anchor','').capitalize()}\s+{int(bloco.get('numero_anchor',0))}\b"
   alvo=None
   for el in elems[i1+1:i2]:
    if isinstance(el, CT_P):
     tx=''.join(el.itertext()).strip()
     if re.match(padrao, tx, re.IGNORECASE): alvo=next((p for p in doc.paragraphs if p._element is el), None); break
   if alvo and bloco.get('texto','').strip(): insert_before(doc, alvo, bloco['texto'].strip(), STYLE_DISC)
 else:
  # fallback: limpa apenas texto antigo antes da 1ª legenda
  first_anchor=None
  for el in elems[i1+1:i2]:
   if isinstance(el, CT_P) and re.match(r'^(Tabela|Figura|Quadro)\s+\d+', ''.join(el.itertext()).strip(), re.IGNORECASE):
    first_anchor=el; break
  for el in elems[i1+1:i2]:
   if el is first_anchor: break
   if isinstance(el, CT_P): el.getparent().remove(el)
  anchor_p=next((p for p in doc.paragraphs if p._element is first_anchor), cons)
  for b in reversed(blocos(texto_md)): insert_before(doc, anchor_p, b, STYLE_DISC)

def substituir_consideracoes(doc, texto):
 t=get_title(doc,'CONSIDERAÇÕES'); r=get_title(doc,'REFERÊNCIAS BIBLIOGRÁFICAS'); ensure_style(doc, STYLE_CONS)
 body=doc._body._element; elems=list(body); i1=elems.index(t._element); i2=elems.index(r._element)
 for el in elems[i1+1:i2]: el.getparent().remove(el)
 bs=blocos(texto)
 if not bs or bs[0]!=ABERTURA_CONSIDERACOES: bs=[ABERTURA_CONSIDERACOES]+([""] if bs else [])+bs
 for b in reversed([x for x in bs if x.strip()]): insert_before(doc, r, b, STYLE_CONS)

def main():
 a=parse_args()
 try:
  doc=Document(a.docx_original); substituir_resultados(doc, carregar(a.resultados), a.blocos_resultados if a.blocos_resultados.exists() else None); substituir_consideracoes(doc, carregar(a.consideracoes)); a.saida_dir.mkdir(parents=True, exist_ok=True); out=a.saida_dir/f"{a.docx_original.stem}_final.docx"; doc.save(out)
 except Exception as e:
  print(f"Erro ao processar DOCX: {e}", file=sys.stderr); return 1
 print(f"Relatório final salvo em: {out}"); return 0
if __name__=='__main__': raise SystemExit(main())
